from tkinter import *
import os
import tkinter.ttk as ttk
import tkinter as tk
from pyperclip import copy, paste
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor

def temizle():
    pano.delete('1.0', END)
    

def yapistir():
    metin=paste()
    
    pano.insert(END,metin)

def hecele():
    hece=pano.get(1.0,END)

    hecele = []

    bits = ''.join(['1' if l in 'AEIİOÖUÜaeıioöuü' else '0' for l in hece])

    seperators = (('101', 1),('1001', 2),('10001', 3))

    index, cut_start_pos = 0, 0

    while index < len(bits):
        for seperator_pattern, seperator_cut_pos in seperators:
            if bits[index:].startswith(seperator_pattern):
                hecele.append(hece[cut_start_pos:index + seperator_cut_pos])

                index += seperator_cut_pos
                cut_start_pos = index
                break

        index += 1

    hecele.append(hece[cut_start_pos:])

    s=len(hecele)     
        
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'TTKB DikTemel Abece'
    font.size = Pt(28)

    paragraph = document.add_paragraph()
    for enum, kelime in enumerate(hecele):
        hece = paragraph.add_run(kelime)
        font = hece.font
        if enum%2==0:
            font.color.rgb = RGBColor(255, 0, 0)
        else:
            font.color.rgb = RGBColor(0, 0, 255)
        
    document.save('heceleme.docx')

    os.startfile("heceleme.docx")

pencere=Tk()
pencere.tk_setPalette("light blue")
pencere.attributes("-fullscreen", 1)

mainframe = ttk.Frame(pencere,padding='3 3 12 12')
mainframe.grid(column=0, row=0)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(0, weight =1)
                    
label=Label(text="HECELEME PROGRAMI",fg="red",font=("TTKB DikTemel Abece" ,20))
label.place(relx = 0.45, rely = 0.0)                  
    
pano=Text( height=10, width=60,fg="black",bg="white",font=("TTKB DikTemel Abece",28))
pano.place(relx = 0.05, rely = 0.1)

buton2=Button()
buton2.config(text="KOPYALADIĞIN METNİ YAPIŞTIR", command=yapistir,width='30',bg="red",fg="white",font=("TTKB DikTemel Abece" ,16))
buton2.place(relx = 0.05, rely = 0.85)

buton1=Button()
buton1.config(text="HECELE", command=hecele,width='25',bg="red",fg="white",font=("TTKB DikTemel Abece" ,16))
buton1.place(relx = 0.30, rely = 0.85)

buton2=Button()
buton2.config(text="TEMİZLE", command=temizle,width='25',bg="red",fg="white",font=("TTKB DikTemel Abece" ,16))
buton2.place(relx = 0.52, rely = 0.85)

buton=Button()
buton.config(text="ÇIKIŞ",command=pencere.destroy,width='25',bg="red",fg="white",font=('TTKB DikTemel Abece',16))
buton.place(relx = 0.74, rely = 0.85)

pencere.mainloop()

